"""Extract all 365 reflections + prologues + intro from DOCX."""
import json, re
from docx import Document

doc = Document('/home/z/my-project/upload/LIBRO EDITADO.docx')

# Phase 1: Find where each section starts by paragraph index
sections = {}
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip().lower()
    if text in ('prólogo 1.', 'prólogo 1', 'prologo 1.', 'prologo 1'):
        sections['prologo1_start'] = i
    elif text in ('prólogo 2.', 'prólogo 2', 'prologo 2.', 'prologo 2'):
        sections['prologo2_start'] = i
    elif text in ('introducción.', 'introducción', 'introduccion.', 'introduccion'):
        sections['intro_start'] = i
    elif text == '1' and i > 60:  # Reflection #1 starts
        sections['reflections_start'] = i
        break

print(f'Section boundaries: {sections}')

# Phase 2: Extract prologue 1 text (between prologo1_start and prologo2_start)
prologo1_text = []
if 'prologo1_start' in sections and 'prologo2_start' in sections:
    for i in range(sections['prologo1_start'] + 1, sections['prologo2_start']):
        t = doc.paragraphs[i].text.strip()
        if t and t.lower() not in ('pastor josé arturo esteves', 'pastor jose arturo esteves'):
            prologo1_text.append(t)

# Phase 3: Extract prologue 2 text (between prologo2_start and intro_start)
prologo2_text = []
if 'prologo2_start' in sections and 'intro_start' in sections:
    for i in range(sections['prologo2_start'] + 1, sections['intro_start']):
        t = doc.paragraphs[i].text.strip()
        if t and t.lower() not in ('julia muñoz de lópez', 'julia muñoz de lopez'):
            prologo2_text.append(t)

# Phase 4: Extract introduction (between intro_start and reflections_start)
intro_text = []
if 'intro_start' in sections and 'reflections_start' in sections:
    for i in range(sections['intro_start'] + 1, sections['reflections_start']):
        t = doc.paragraphs[i].text.strip()
        if t:
            intro_text.append(t)

print(f'Prologo 1: {len(prologo1_text)} paragraphs')
print(f'Prologo 2: {len(prologo2_text)} paragraphs')
print(f'Introduccion: {len(intro_text)} paragraphs')

# Phase 5: Extract all reflections starting from reflections_start
reflexiones = []
current_ref = None
body_lines = []

start = sections.get('reflections_start', 85)
for i in range(start, len(doc.paragraphs)):
    text = doc.paragraphs[i].text.strip()
    if not text:
        continue

    # Check if this is a reflection number
    if text.isdigit():
        num = int(text)
        if 1 <= num <= 365:
            # Save previous reflection
            if current_ref is not None:
                current_ref['body'] = ' '.join(body_lines).strip()
                reflexiones.append(current_ref)
            current_ref = {'number': num, 'title': '', 'quote': '', 'body': ''}
            body_lines = []
            continue

    # If we're in a reflection
    if current_ref is not None:
        if not current_ref['title'] and not text.startswith('"'):
            current_ref['title'] = text
        elif text.startswith('"'):
            if not current_ref['quote']:
                current_ref['quote'] = text.strip('"')
            else:
                body_lines.append(text)
        else:
            body_lines.append(text)

# Save last reflection
if current_ref is not None:
    current_ref['body'] = ' '.join(body_lines).strip()
    reflexiones.append(current_ref)

print(f'Reflexiones extraidas: {len(reflexiones)}')
if reflexiones:
    print(f'Primera: #{reflexiones[0]["number"]} - {reflexiones[0]["title"]}')
    print(f'Ultima: #{reflexiones[-1]["number"]} - {reflexiones[-1]["title"]}')
    with_q = sum(1 for r in reflexiones if r['quote'])
    print(f'Con cita: {with_q}')
    nums = {r['number'] for r in reflexiones}
    missing = [n for n in range(1, 366) if n not in nums]
    if missing:
        print(f'Faltantes: {missing}')
    else:
        print('¡Todas las 365 reflexiones presentes!')

    # Show sample
    r = reflexiones[0]
    print(f'\nCita ejemplo: {r["quote"][:150]}...')
    print(f'Body largo: {len(r["body"])} chars')

# Build output
output = {
    'prologo1': {
        'title': 'Prólogo',
        'author': 'Pastor José Arturo Esteves',
        'text': ' '.join(prologo1_text)
    },
    'prologo2': {
        'title': 'Prólogo',
        'author': 'Julia Muñoz de López',
        'text': ' '.join(prologo2_text)
    },
    'introduccion': {
        'title': 'Introducción',
        'author': 'Pastor Nicolás Abreu',
        'text': ' '.join(intro_text)
    },
    'reflexiones': reflexiones
}

with open('/home/z/my-project/src/data/reflexiones.json', 'w', encoding='utf-8') as f:
    json.dump(output, f, ensure_ascii=False, indent=2)
print(f'\nGuardado en src/data/reflexiones.json ({sum(1 for _ in reflexiones)} reflexiones)')
